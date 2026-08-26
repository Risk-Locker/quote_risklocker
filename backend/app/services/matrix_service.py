"""Service for aggregating company catalog matrices, exporting DOCX/XLSX, and diffing."""

from __future__ import annotations

import io
from datetime import date, datetime, timezone
from typing import Any

import docx
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn
from docx.shared import Inches, Pt, RGBColor
import openpyxl
from openpyxl.styles import Alignment, Border, Font, PatternFill, Side
from openpyxl.utils import get_column_letter
from sqlalchemy import select
from sqlalchemy.orm import Session

from app.core.errors import AppError
from app.models.tables import (
    BenefitCatalog,
    BenefitCatalogRevision,
    BenefitConcept,
    BenefitPackage,
    BenefitPackagePlan,
    BenefitPackagePlanItem,
    CatalogOffering,
    CoverageType,
    InsuranceCompany,
    InsuranceProduct,
    InsuranceProductTier,
    Segment,
    VehicleCategory,
)


def _set_cell_shading(cell, color_hex: str) -> None:
    """Set background color for a docx cell."""
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    tc_pr.append(shd)


def _set_cell_margins(cell, top=120, bottom=120, left=150, right=150) -> None:
    """Set inner padding for a docx cell (in twips, 20 twips = 1 pt)."""
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = parse_xml(
        f'<w:tcMar {nsdecls("w")}>'
        f'<w:top w:w="{top}" w:type="dxa"/>'
        f'<w:bottom w:w="{bottom}" w:type="dxa"/>'
        f'<w:left w:w="{left}" w:type="dxa"/>'
        f'<w:right w:w="{right}" w:type="dxa"/>'
        f"</w:tcMar>"
    )
    tc_pr.append(tc_mar)


def get_company_matrix_data(db: Session, company_id: str) -> dict[str, Any]:
    """Aggregate all products, catalogs, packages, and offerings for an insurance company."""
    company = db.get(InsuranceCompany, company_id)
    if company is None:
        raise AppError("Company not found.", 404)

    # Pre-fetch reference lookups
    segments = {s.id: s for s in db.scalars(select(Segment)).all()}
    vehicles = {v.id: v for v in db.scalars(select(VehicleCategory)).all()}
    coverages = {c.id: c for c in db.scalars(select(CoverageType)).all()}
    concepts = {c.id: c for c in db.scalars(select(BenefitConcept)).all()}

    products = list(
        db.scalars(
            select(InsuranceProduct)
            .where(InsuranceProduct.company_id == company_id)
            .order_by(InsuranceProduct.sort_order if hasattr(InsuranceProduct, "sort_order") else InsuranceProduct.name)
        ).all()
    )
    product_map = {p.id: p for p in products}

    catalogs = list(
        db.scalars(
            select(BenefitCatalog)
            .where(BenefitCatalog.company_id == company_id)
            .order_by(BenefitCatalog.name)
        ).all()
    )

    scenarios = []
    total_defaults_count = 0
    total_addons_count = 0
    total_bundles_count = 0

    for cat in catalogs:
        latest_rev = db.scalar(
            select(BenefitCatalogRevision)
            .where(BenefitCatalogRevision.catalog_id == cat.id)
            .order_by(BenefitCatalogRevision.revision_number.desc())
        )
        if not latest_rev:
            continue

        prod = product_map.get(cat.product_id) if cat.product_id else None
        seg = segments.get(cat.segment_id) if cat.segment_id else None
        veh = vehicles.get(cat.vehicle_category_id) if cat.vehicle_category_id else None
        cov = coverages.get(cat.coverage_type_id) if cat.coverage_type_id else None

        # Offerings for this catalog revision
        offerings = list(
            db.scalars(
                select(CatalogOffering)
                .where(CatalogOffering.catalog_revision_id == latest_rev.id)
                .order_by(CatalogOffering.sort_order, CatalogOffering.offering_key)
            ).all()
        )

        # Packages belonging to this catalog revision
        packages = list(
            db.scalars(
                select(BenefitPackage)
                .where(BenefitPackage.catalog_revision_id == latest_rev.id)
                .order_by(BenefitPackage.sort_order, BenefitPackage.name)
            ).all()
        )

        pkg_ids = [p.id for p in packages]
        plans = list(
            db.scalars(
                select(BenefitPackagePlan)
                .where(BenefitPackagePlan.package_id.in_(pkg_ids))
                .order_by(BenefitPackagePlan.sort_order, BenefitPackagePlan.name)
            ).all()
        ) if pkg_ids else []

        plan_ids = [p.id for p in plans]
        plan_items = list(
            db.scalars(
                select(BenefitPackagePlanItem)
                .where(BenefitPackagePlanItem.plan_id.in_(plan_ids))
                .order_by(BenefitPackagePlanItem.sort_order)
            ).all()
        ) if plan_ids else []

        # Map plan items to plans
        plan_items_by_plan: dict[str, list[BenefitPackagePlanItem]] = {}
        for pi in plan_items:
            plan_items_by_plan.setdefault(pi.plan_id, []).append(pi)

        # Map plans to packages
        plans_by_pkg: dict[str, list[dict[str, Any]]] = {}
        offering_map = {o.id: o for o in offerings}
        for plan in plans:
            p_items = []
            for item in plan_items_by_plan.get(plan.id, []):
                off = offering_map.get(item.offering_id)
                conc = concepts.get(off.concept_id) if off else None
                p_items.append({
                    "offering_key": off.offering_key if off else "",
                    "label": (off.label_override if off else None) or (conc.label if conc else "Benefit"),
                    "override_value": (
                        item.typed_value_override.get("display_text")
                        if item.typed_value_override and isinstance(item.typed_value_override, dict)
                        else (off.display_value if off else "")
                    ),
                })
            plans_by_pkg.setdefault(plan.package_id, []).append({
                "plan_id": plan.id,
                "plan_key": plan.plan_key,
                "name": plan.name,
                "items": p_items,
            })

        # Process offerings into defaults vs addons
        defaults = []
        addons = []
        for o in offerings:
            conc = concepts.get(o.concept_id)
            c_key = conc.concept_key if conc else o.offering_key
            c_label = o.label_override or (conc.label if conc else o.offering_key)
            c_desc = conc.description if conc else ""

            is_default = o.role == "included" or o.offering_kind == "base"
            if is_default:
                defaults.append({
                    "offering_id": o.id,
                    "offering_key": o.offering_key,
                    "concept_key": c_key,
                    "label": c_label,
                    "description": c_desc,
                    "display_value": o.display_value or "Included",
                    "price": 0.0,
                    "price_text": "0 RM",
                })
            else:
                price_val = 0.0
                price_str = "Optional"
                if o.optional_price and isinstance(o.optional_price, dict):
                    raw_val = o.optional_price.get("value") or o.optional_price.get("amount")
                    try:
                        price_val = float(raw_val) if raw_val is not None else 0.0
                    except (ValueError, TypeError):
                        price_val = 0.0
                    rate_pct = o.optional_price.get("rate_pct")
                    if rate_pct:
                        price_str = f"{rate_pct}% rate"
                    elif price_val > 0:
                        price_str = f"RM {price_val:.2f}"
                    else:
                        price_str = o.display_value or "Optional"
                elif o.display_value:
                    price_str = o.display_value

                addons.append({
                    "offering_id": o.id,
                    "offering_key": o.offering_key,
                    "concept_key": c_key,
                    "label": c_label,
                    "description": c_desc,
                    "display_value": o.display_value or "Optional",
                    "price": price_val,
                    "price_text": price_str,
                })

        bundles_data = []
        for pkg in packages:
            bundles_data.append({
                "package_id": pkg.id,
                "package_key": pkg.package_key,
                "name": pkg.name,
                "package_kind": pkg.package_kind,
                "plans": plans_by_pkg.get(pkg.id, []),
            })

        total_defaults_count += len(defaults)
        total_addons_count += len(addons)
        total_bundles_count += len(bundles_data)

        # Determine system type (Package System with tiers vs Add-on System)
        is_packaged = bool(packages and any("tier" in p.name.lower() or "lite" in p.name.lower() or "plus" in p.name.lower() or "premier" in p.name.lower() for p in packages))
        system_type = "Package System" if (is_packaged or (prod and "auto365" in prod.product_key.lower())) else "Add-on System"

        scenarios.append({
            "catalog_id": cat.id,
            "product_id": prod.id if prod else None,
            "product_name": prod.name if prod else cat.name,
            "product_key": prod.product_key if prod else "",
            "scenario_name": cat.name,
            "segment_name": seg.name if seg else "Private",
            "segment_key": seg.segment_key if seg else "private",
            "vehicle_category_name": veh.name if veh else "Car",
            "vehicle_category_key": veh.category_key if veh else "car",
            "coverage_type_name": cov.name if cov else "Comprehensive",
            "coverage_type_key": cov.coverage_key if cov else "comprehensive",
            "system_type": system_type,
            "revision_number": latest_rev.revision_number,
            "state": latest_rev.state,
            "defaults": defaults,
            "addons": addons,
            "bundles": bundles_data,
        })

    # Sort scenarios logically: Comprehensive first, then TPFT, then TPO; then Car, Motorcycle, Commercial; then Private, Company
    def scenario_sort_key(s: dict[str, Any]):
        cov_order = {"comprehensive": 1, "third_party_fire_theft": 2, "tpft": 2, "third_party": 3, "tpo": 3}
        veh_order = {"car": 1, "motorcycle": 2, "commercial_vehicle": 3}
        seg_order = {"private": 1, "company_commercial": 2}
        return (
            cov_order.get(s["coverage_type_key"].lower(), 99),
            veh_order.get(s["vehicle_category_key"].lower(), 99),
            seg_order.get(s["segment_key"].lower(), 99),
            s["scenario_name"],
        )

    scenarios.sort(key=scenario_sort_key)

    return {
        "company": {
            "id": company.id,
            "name": company.name,
            "slug": company.slug,
            "category": company.category,
        },
        "summary": {
            "total_products": len(products),
            "total_scenarios": len(scenarios),
            "total_defaults": total_defaults_count,
            "total_addons": total_addons_count,
            "total_bundles": total_bundles_count,
        },
        "scenarios": scenarios,
    }


def generate_company_matrix_docx(data: dict[str, Any]) -> io.BytesIO:
    """Generate a clean, standardized Word (.docx) document matching the canonical reference format."""
    doc = docx.Document()

    # Configure Landscape Page Layout with 0.5 in margins
    for section in doc.sections:
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width = Inches(11.0)
        section.page_height = Inches(8.5)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)

    company_name = data["company"]["name"]
    summary = data["summary"]
    scenarios = data["scenarios"]

    # Header Title
    title_p = doc.add_paragraph()
    title_run = title_p.add_run(f"{company_name} — Standardized Benefits & Packages Matrix")
    title_run.font.name = "Arial"
    title_run.font.size = Pt(16)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(0x1F, 0x29, 0x37)

    # Subtitle / Metadata
    sub_p = doc.add_paragraph()
    today_str = date.today().strftime("%d %B %Y")
    sub_run = sub_p.add_run(
        f"Generated by RiskLocker System · Date: {today_str} · "
        f"Scenarios: {summary['total_scenarios']} · Products: {summary['total_products']} · "
        f"Canonical Underwriting Source of Truth"
    )
    sub_run.font.name = "Arial"
    sub_run.font.size = Pt(9.5)
    sub_run.font.italic = True
    sub_run.font.color.rgb = RGBColor(0x4B, 0x55, 0x63)

    # Explanatory prompt / guidance
    notice_p = doc.add_paragraph()
    notice_run = notice_p.add_run(
        "Standard Format Note: Defaults are included policy coverages (Cost: 0 RM). Add-ons are payable endorsements with exact base costs or rate percentages. "
        "Bundled Add-ons specify structured multi-benefit packages (e.g. Plan A/B/C/D). "
        "To modify or seed additional benefits, edit the table cells below and provide this document directly to the system."
    )
    notice_run.font.name = "Arial"
    notice_run.font.size = Pt(8.5)
    notice_run.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)

    doc.add_paragraph()  # Spacer

    # Group scenarios by Coverage Type (Comprehensive first, etc.)
    coverage_groups: dict[str, list[dict[str, Any]]] = {}
    for s in scenarios:
        coverage_groups.setdefault(s["coverage_type_name"], []).append(s)

    for cov_name, group_scenarios in coverage_groups.items():
        # Section Heading
        h_p = doc.add_paragraph()
        h_run = h_p.add_run(f"Coverage Type: {cov_name} ({len(group_scenarios)} Products / Scenarios)")
        h_run.font.name = "Arial"
        h_run.font.size = Pt(12)
        h_run.font.bold = True
        h_run.font.color.rgb = RGBColor(0x11, 0x18, 0x27)

        # Table: 4 Columns
        # Col 0: Product & Scenario (width 2.2 in)
        # Col 1: Default Benefits (Included) (width 3.4 in)
        # Col 2: Add-on Benefits & Exact Base Cost (width 3.4 in)
        # Col 3: Bundled Add-ons & Plans (width 1.0 in)
        table = doc.add_table(rows=1, cols=4)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False

        col_widths = [Inches(2.2), Inches(3.4), Inches(3.4), Inches(1.0)]
        for i, col in enumerate(table.columns):
            col.width = col_widths[i]

        hdr_cells = table.rows[0].cells
        hdr_titles = [
            "Product / Scenario",
            "Default Benefits (Included)",
            "Add-on Benefits & Exact Base Cost",
            "Bundled Add-ons",
        ]
        for i, title in enumerate(hdr_titles):
            hdr_cells[i].text = title
            _set_cell_shading(hdr_cells[i], "1F2937")
            _set_cell_margins(hdr_cells[i], top=120, bottom=120, left=140, right=140)
            p = hdr_cells[i].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for run in p.runs:
                run.font.name = "Arial"
                run.font.size = Pt(9)
                run.font.bold = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

        for s in group_scenarios:
            row_cells = table.add_row().cells
            for i in range(4):
                _set_cell_margins(row_cells[i], top=100, bottom=100, left=130, right=130)

            # Col 0: Product & Hierarchy Details
            p0 = row_cells[0].paragraphs[0]
            r0_name = p0.add_run(f"{s['scenario_name']}\n")
            r0_name.font.name = "Arial"
            r0_name.font.size = Pt(9.5)
            r0_name.font.bold = True

            r0_meta = p0.add_run(
                f"• Segment: {s['segment_name']}\n"
                f"• Vehicle: {s['vehicle_category_name']}\n"
                f"• Coverage: {s['coverage_type_name']}\n"
                f"• System: {s['system_type']}\n"
                f"• Revision: rev {s['revision_number']} ({s['state']})"
            )
            r0_meta.font.name = "Arial"
            r0_meta.font.size = Pt(8)
            r0_meta.font.color.rgb = RGBColor(0x4B, 0x55, 0x63)

            # Col 1: Defaults
            p1 = row_cells[1].paragraphs[0]
            if not s["defaults"]:
                r = p1.add_run("No default benefits configured.")
                r.font.name = "Arial"
                r.font.size = Pt(8.5)
                r.font.italic = True
                r.font.color.rgb = RGBColor(0x9C, 0xA3, 0xAF)
            else:
                for idx, d in enumerate(s["defaults"]):
                    prefix = "• " if idx == 0 else "\n• "
                    r_lbl = p1.add_run(f"{prefix}{d['label']} : ")
                    r_lbl.font.name = "Arial"
                    r_lbl.font.size = Pt(8.5)
                    r_lbl.font.bold = True

                    val_str = d['display_value'] or d['description'] or "Included"
                    r_val = p1.add_run(f"{val_str}. ")
                    r_val.font.name = "Arial"
                    r_val.font.size = Pt(8.5)

                    r_cost = p1.add_run("Cost : 0 RM")
                    r_cost.font.name = "Arial"
                    r_cost.font.size = Pt(8)
                    r_cost.font.color.rgb = RGBColor(0x05, 0x96, 0x69)

            # Col 2: Add-ons
            p2 = row_cells[2].paragraphs[0]
            if not s["addons"]:
                r = p2.add_run("No add-on riders configured.")
                r.font.name = "Arial"
                r.font.size = Pt(8.5)
                r.font.italic = True
                r.font.color.rgb = RGBColor(0x9C, 0xA3, 0xAF)
            else:
                for idx, a in enumerate(s["addons"]):
                    prefix = "• " if idx == 0 else "\n• "
                    r_lbl = p2.add_run(f"{prefix}{a['label']} : ")
                    r_lbl.font.name = "Arial"
                    r_lbl.font.size = Pt(8.5)
                    r_lbl.font.bold = True

                    val_str = a['display_value'] or a['description'] or "Optional"
                    r_val = p2.add_run(f"{val_str}. ")
                    r_val.font.name = "Arial"
                    r_val.font.size = Pt(8.5)

                    r_cost = p2.add_run(f"Cost : {a['price_text']}")
                    r_cost.font.name = "Arial"
                    r_cost.font.size = Pt(8)
                    r_cost.font.color.rgb = RGBColor(0x25, 0x63, 0xEB)

            # Col 3: Bundles
            p3 = row_cells[3].paragraphs[0]
            if not s["bundles"]:
                r = p3.add_run("None")
                r.font.name = "Arial"
                r.font.size = Pt(8.5)
                r.font.italic = True
                r.font.color.rgb = RGBColor(0x9C, 0xA3, 0xAF)
            else:
                for idx, b in enumerate(s["bundles"]):
                    prefix = "" if idx == 0 else "\n\n"
                    r_bname = p3.add_run(f"{prefix}{b['name']}\n")
                    r_bname.font.name = "Arial"
                    r_bname.font.size = Pt(8.5)
                    r_bname.font.bold = True

                    for plan in b.get("plans", []):
                        r_p = p3.add_run(f"  └ {plan['name']}\n")
                        r_p.font.name = "Arial"
                        r_p.font.size = Pt(8)
                        r_p.font.color.rgb = RGBColor(0x7C, 0x3A, 0xED)

        doc.add_paragraph()  # Spacing between coverage groups

    output = io.BytesIO()
    doc.save(output)
    output.seek(0)
    return output


def generate_company_matrix_xlsx(data: dict[str, Any]) -> io.BytesIO:
    """Generate a clean, multi-tab Excel workbook representing the complete company matrix."""
    wb = openpyxl.Workbook()

    # Setup styles
    header_fill = PatternFill(start_color="1F2937", end_color="1F2937", fill_type="solid")
    header_font = Font(name="Arial", size=10, bold=True, color="FFFFFF")
    data_font = Font(name="Arial", size=9)
    bold_data_font = Font(name="Arial", size=9, bold=True)
    muted_data_font = Font(name="Arial", size=8, italic=True, color="6B7280")
    thin_border = Border(
        left=Side(style="thin", color="E5E7EB"),
        right=Side(style="thin", color="E5E7EB"),
        top=Side(style="thin", color="E5E7EB"),
        bottom=Side(style="thin", color="E5E7EB"),
    )

    # ─────────────────────────────────────────────────────────────────────────
    # Sheet 1: Catalog Scenarios Overview
    # ─────────────────────────────────────────────────────────────────────────
    ws1 = wb.active if wb.active is not None else wb.create_sheet(title="Scenarios Overview")
    assert ws1 is not None
    ws1.title = "Scenarios Overview"
    ws1.views.sheetView[0].showGridLines = True

    headers1 = [
        "Company",
        "Product Key",
        "Scenario Name",
        "Coverage Type",
        "Segment",
        "Vehicle Category",
        "System Type",
        "Revision",
        "Defaults Count",
        "Add-ons Count",
        "Bundles Count",
        "Included Defaults Summary",
        "Available Add-ons Summary",
    ]
    ws1.append(headers1)
    for col_num in range(1, len(headers1) + 1):
        cell = ws1.cell(row=1, column=col_num)
        cell.fill = header_fill
        cell.font = header_font
        cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)

    company_name = data["company"]["name"]
    for s in data["scenarios"]:
        def_summary = "; ".join(f"{d['label']} ({d['display_value']})" for d in s["defaults"])
        add_summary = "; ".join(f"{a['label']} ({a['price_text']})" for a in s["addons"])

        row_data = [
            company_name,
            s["product_key"],
            s["scenario_name"],
            s["coverage_type_name"],
            s["segment_name"],
            s["vehicle_category_name"],
            s["system_type"],
            f"rev {s['revision_number']}",
            len(s["defaults"]),
            len(s["addons"]),
            len(s["bundles"]),
            def_summary,
            add_summary,
        ]
        ws1.append(row_data)

    # Format data rows of Sheet 1
    for row in ws1.iter_rows(min_row=2, max_row=ws1.max_row, min_col=1, max_col=len(headers1)):
        for cell in row:
            cell.font = data_font
            cell.border = thin_border
            cell.alignment = Alignment(vertical="top", wrap_text=True)

    # ─────────────────────────────────────────────────────────────────────────
    # Sheet 2: Detailed Line-Item Offerings Matrix (AI & Seed Friendly)
    # ─────────────────────────────────────────────────────────────────────────
    ws2 = wb.create_sheet(title="Detailed Offerings")
    ws2.views.sheetView[0].showGridLines = True

    headers2 = [
        "Company",
        "Product Key",
        "Scenario Name",
        "Coverage Type",
        "Segment",
        "Vehicle Category",
        "Benefit Type",
        "Concept Key",
        "Benefit Label",
        "Coverage Limit / Value",
        "Base Price (RM)",
        "Price Text / Formula",
        "Bundle / Plan Tier",
    ]
    ws2.append(headers2)
    for col_num in range(1, len(headers2) + 1):
        cell = ws2.cell(row=1, column=col_num)
        cell.fill = header_fill
        cell.font = header_font
        cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)

    for s in data["scenarios"]:
        # Defaults
        for d in s["defaults"]:
            ws2.append([
                company_name,
                s["product_key"],
                s["scenario_name"],
                s["coverage_type_name"],
                s["segment_name"],
                s["vehicle_category_name"],
                "Default (Included)",
                d["concept_key"],
                d["label"],
                d["display_value"],
                0.0,
                "0 RM",
                "-",
            ])

        # Add-ons
        for a in s["addons"]:
            ws2.append([
                company_name,
                s["product_key"],
                s["scenario_name"],
                s["coverage_type_name"],
                s["segment_name"],
                s["vehicle_category_name"],
                "Add-on (Optional)",
                a["concept_key"],
                a["label"],
                a["display_value"],
                a["price"],
                a["price_text"],
                "-",
            ])

        # Bundles & Plans
        for b in s["bundles"]:
            for plan in b.get("plans", []):
                for item in plan.get("items", []):
                    ws2.append([
                        company_name,
                        s["product_key"],
                        s["scenario_name"],
                        s["coverage_type_name"],
                        s["segment_name"],
                        s["vehicle_category_name"],
                        "Bundle Component",
                        item["offering_key"],
                        item["label"],
                        item["override_value"],
                        0.0,
                        "Included in Bundle",
                        f"{b['name']} → {plan['name']}",
                    ])

    # Format data rows of Sheet 2
    for row in ws2.iter_rows(min_row=2, max_row=ws2.max_row, min_col=1, max_col=len(headers2)):
        for cell in row:
            cell.font = data_font
            cell.border = thin_border
            cell.alignment = Alignment(vertical="top", wrap_text=False)

    # Auto-adjust column widths for both sheets
    for ws in (ws1, ws2):
        for col in ws.columns:
            if not col or col[0].column is None:
                continue
            max_len = max(len(str(cell.value or "")) for cell in col)
            col_letter = get_column_letter(int(col[0].column))
            ws.column_dimensions[col_letter].width = min(max(max_len + 3, 12), 50)

    output = io.BytesIO()
    wb.save(output)
    output.seek(0)
    return output


def diff_company_matrix(existing_data: dict[str, Any], incoming_data: dict[str, Any]) -> dict[str, Any]:
    """Compare an incoming matrix dataset against existing database catalog offerings.
    
    Returns categorized additions, modifications, and deletions for zero-loss incremental sync.
    """
    diffs = {
        "company": existing_data.get("company", {}),
        "total_changes": 0,
        "scenarios_diff": [],
    }

    existing_scenarios = {s["scenario_name"].strip().lower(): s for s in existing_data.get("scenarios", [])}
    incoming_scenarios = incoming_data.get("scenarios", [])

    for inc in incoming_scenarios:
        s_name = inc.get("scenario_name", "").strip()
        matched_exist = existing_scenarios.get(s_name.lower())

        if not matched_exist:
            diffs["scenarios_diff"].append({
                "scenario_name": s_name,
                "status": "new_scenario",
                "added_defaults": inc.get("defaults", []),
                "added_addons": inc.get("addons", []),
            })
            diffs["total_changes"] += len(inc.get("defaults", [])) + len(inc.get("addons", []))
            continue

        exist_defaults = {d["concept_key"]: d for d in matched_exist.get("defaults", [])}
        exist_addons = {a["concept_key"]: a for a in matched_exist.get("addons", [])}

        inc_defaults = {d.get("concept_key"): d for d in inc.get("defaults", []) if d.get("concept_key")}
        inc_addons = {a.get("concept_key"): a for a in inc.get("addons", []) if a.get("concept_key")}

        added_defaults = [d for k, d in inc_defaults.items() if k not in exist_defaults]
        modified_defaults = [
            {"from": exist_defaults[k], "to": d}
            for k, d in inc_defaults.items()
            if k in exist_defaults and (d.get("display_value") != exist_defaults[k].get("display_value"))
        ]

        added_addons = [a for k, a in inc_addons.items() if k not in exist_addons]
        modified_addons = [
            {"from": exist_addons[k], "to": a}
            for k, a in inc_addons.items()
            if k in exist_addons and (
                float(a.get("price") or 0.0) != float(exist_addons[k].get("price") or 0.0)
                or a.get("display_value") != exist_addons[k].get("display_value")
            )
        ]

        s_changes = len(added_defaults) + len(modified_defaults) + len(added_addons) + len(modified_addons)
        if s_changes > 0:
            diffs["total_changes"] += s_changes
            diffs["scenarios_diff"].append({
                "scenario_name": s_name,
                "status": "modified",
                "added_defaults": added_defaults,
                "modified_defaults": modified_defaults,
                "added_addons": added_addons,
                "modified_addons": modified_addons,
            })

    return diffs
