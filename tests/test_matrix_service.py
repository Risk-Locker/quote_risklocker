import io
import sys
from pathlib import Path
from unittest.mock import MagicMock

ROOT = Path(__file__).resolve().parents[1]
BACKEND = ROOT / "backend"
if str(BACKEND) not in sys.path:
    sys.path.insert(0, str(BACKEND))

import docx
import openpyxl
from app.services.matrix_service import (
    get_company_matrix_data,
    generate_company_matrix_docx,
    generate_company_matrix_xlsx,
    diff_company_matrix,
)
from app.models.tables import InsuranceCompany


def test_matrix_document_generation():
    test_data = {
        "company": {
            "id": "comp-123",
            "name": "QBE Insurance (Malaysia) Berhad",
            "slug": "qbe",
            "category": "Motor",
        },
        "summary": {
            "total_products": 2,
            "total_scenarios": 2,
            "total_defaults": 4,
            "total_addons": 4,
            "total_bundles": 1,
        },
        "scenarios": [
            {
                "catalog_id": "cat-1",
                "product_id": "prod-1",
                "product_name": "Private Car Protector",
                "product_key": "qbe-private-car",
                "scenario_name": "Private Car Protector",
                "segment_name": "Private",
                "segment_key": "private",
                "vehicle_category_name": "Car",
                "vehicle_category_key": "car",
                "coverage_type_name": "Comprehensive",
                "coverage_type_key": "comprehensive",
                "system_type": "Add-on System",
                "revision_number": 3,
                "state": "published",
                "defaults": [
                    {
                        "offering_id": "off-1",
                        "offering_key": "own-damage",
                        "concept_key": "own-damage",
                        "label": "Own Damage",
                        "description": "Accidental damage",
                        "display_value": "Included",
                        "price": 0.0,
                        "price_text": "0 RM",
                    },
                    {
                        "offering_id": "off-2",
                        "offering_key": "third-party-bi",
                        "concept_key": "third-party-bi",
                        "label": "Third Party Bodily Injury",
                        "description": "Unlimited statutory liability",
                        "display_value": "Unlimited",
                        "price": 0.0,
                        "price_text": "0 RM",
                    },
                ],
                "addons": [
                    {
                        "offering_id": "off-3",
                        "offering_key": "windscreen",
                        "concept_key": "windscreen",
                        "label": "Windscreen Cover",
                        "description": "Windscreen repair up to RM1,000",
                        "display_value": "RM 1,000",
                        "price": 150.0,
                        "price_text": "RM 150.00",
                    },
                    {
                        "offering_id": "off-4",
                        "offering_key": "special-perils",
                        "concept_key": "special-perils",
                        "label": "Special Perils",
                        "description": "Full flood cover",
                        "display_value": "Full Flood",
                        "price": 125.0,
                        "price_text": "0.25% rate",
                    },
                ],
                "bundles": [
                    {
                        "package_id": "pkg-1",
                        "package_key": "dpp",
                        "name": "Driver Passenger Protector",
                        "package_kind": "comprehensive",
                        "plans": [
                            {
                                "plan_id": "plan-1",
                                "plan_key": "plan-a",
                                "name": "Plan A (RM 70)",
                                "items": [
                                    {
                                        "offering_key": "pa",
                                        "label": "PA Death Benefit",
                                        "override_value": "RM 50,000",
                                    }
                                ],
                            }
                        ],
                    }
                ],
            }
        ],
    }

    # 1. DOCX generation
    docx_buf = generate_company_matrix_docx(test_data)
    assert isinstance(docx_buf, io.BytesIO)
    docx_bytes = docx_buf.getvalue()
    assert len(docx_bytes) > 5000
    loaded_doc = docx.Document(io.BytesIO(docx_bytes))
    assert len(loaded_doc.tables) >= 1
    assert "QBE Insurance" in loaded_doc.paragraphs[0].text

    # 2. XLSX generation
    xlsx_buf = generate_company_matrix_xlsx(test_data)
    assert isinstance(xlsx_buf, io.BytesIO)
    xlsx_bytes = xlsx_buf.getvalue()
    assert len(xlsx_bytes) > 3000
    loaded_wb = openpyxl.load_workbook(io.BytesIO(xlsx_bytes))
    assert "Scenarios Overview" in loaded_wb.sheetnames
    assert "Detailed Offerings" in loaded_wb.sheetnames
    ws2 = loaded_wb["Detailed Offerings"]
    assert ws2.max_row >= 4

    # 3. Diffing logic
    incoming = {
        "scenarios": [
            {
                "scenario_name": "Private Car Protector",
                "defaults": [
                    {"concept_key": "own-damage", "display_value": "Included"},
                    {"concept_key": "new-default", "display_value": "Free Towing 100km"},
                ],
                "addons": [
                    {"concept_key": "windscreen", "display_value": "RM 1,000", "price": 160.0},
                    {"concept_key": "brand-new-addon", "display_value": "New Cover", "price": 50.0},
                ],
            }
        ]
    }
    diff = diff_company_matrix(test_data, incoming)
    assert diff["total_changes"] == 3
    assert len(diff["scenarios_diff"]) == 1
    s_diff = diff["scenarios_diff"][0]
    assert len(s_diff["added_defaults"]) == 1
    assert s_diff["added_defaults"][0]["concept_key"] == "new-default"
    assert len(s_diff["modified_addons"]) == 1
    assert s_diff["modified_addons"][0]["to"]["price"] == 160.0
    assert len(s_diff["added_addons"]) == 1
    assert s_diff["added_addons"][0]["concept_key"] == "brand-new-addon"
